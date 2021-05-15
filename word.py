import docx

d = docx.Document('demo.docx')

print(d.paragraphs[0].text)

p = d.paragraphs[1]

print(p.runs[0].text)

print(p.runs[1].bold)
print(p.runs[1].bold)
print(p.runs[3].italic)
print(p.runs[2].underline)
p.runs[2].underline == True
print(p.runs[2].underline)
p.runs[3].text == 'italic :)'

d.save('demo2.docx')

p.style = 'Title'
d.save('demo3.docx')


"""
Nowy dokument
"""

d = docx.Document()
d.add_paragraph('This is a paragraph.')
d.add_paragraph('Andathor one.')
d.save('demo4.docx')

"""
całość dokumentu jako string
"""

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for p in doc.paragraphs:
        fullText.append((p.text))
    return '\n'.join(fullText)


print(getText('demo.docx'))
